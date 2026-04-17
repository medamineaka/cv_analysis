import docx
import pandas as pd

import pytesseract
from pdf2image import convert_from_path

import os

def obtenir_blocs_docx(chemin_docx):
    doc = docx.Document(chemin_docx)
    tous_les_blocs = []

    # 1. Extraction des paragraphes (avec détection de style)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() != "":
            tous_les_blocs.append({
                'source': 'docx_para',
                'ordre': i,
                'texte': para.text.strip(),
                'style': para.style.name,  # Très important pour l'IA (Titre, Normal, etc.)
                'is_bold': any(run.bold for run in para.runs)
            })

    # 2. Extraction des tableaux (souvent utilisés pour les compétences)
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() != "":
                    tous_les_blocs.append({
                        'source': 'docx_table',
                        'ordre': i,
                        'texte': cell.text.strip(),
                        'style': 'Table_Cell',
                        'is_bold': False
                    })

    return pd.DataFrame(tous_les_blocs)




def obtenir_blocs_pdf(chemin_pdf):
    # 1. Conversion du PDF en images (Haute résolution 300 DPI)
    # Cela permet de traiter le texte "dessiné"
    pages = convert_from_path(chemin_pdf, 300)
    tous_les_blocs = []

    for num_page, image in enumerate(pages):
        # 2. Appel à Tesseract avec le mode "image_to_data"
        # Ce mode renvoie : texte, confiance, x, y, largeur, hauteur
        data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT, lang='fra+eng')

        for i in range(len(data['text'])):
            texte = data['text'][i].strip()
            confiance = int(data['conf'][i])

            # 3. Filtrage : on ne garde que les blocs lisibles (>60%)
            if confiance > 60 and texte != "":
                tous_les_blocs.append({
                    'source': 'pdf_ocr',
                    'page': num_page,
                    'texte': texte,
                    'x': data['left'][i],  # Position horizontale
                    'y': data['top'][i],  # Position verticale
                    'w': data['width'][i],  # Largeur du bloc
                    'h': data['height'][i],  # Hauteur du bloc
                    'conf': confiance
                })

    return pd.DataFrame(tous_les_blocs)





def extracteur_universel(chemin_fichier):
    extension = os.path.splitext(chemin_fichier)[1].lower()

    if extension == '.pdf':
        return obtenir_blocs_pdf(chemin_fichier)
    elif extension == '.docx':
        return obtenir_blocs_docx(chemin_fichier)
    else:
        raise ValueError("Format non supporté (Utilisez PDF ou DOCX)")