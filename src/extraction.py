# extraction.py
"""
Extraction universelle de texte depuis CV (PDF/DOCX) et Offres (JSON texte).
- PDF natif  : pymupdf + KMeans détection colonnes automatique
- PDF scanné : tesseract OCR + KMeans colonnes
- DOCX       : python-docx paragraphes + tableaux
- Offre      : texte brut LinkedIn (20-30 lignes)
"""

import os
import json
import re
import docx
import fitz
import pytesseract
import numpy as np
from pdf2image import convert_from_path
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score


# ════════════════════════════════════════════════
# 1. DOCX
# ════════════════════════════════════════════════

def _extraire_docx(path: str) -> str:
    doc   = docx.Document(path)
    lines = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


# ════════════════════════════════════════════════
# 2. DÉTECTION COLONNES
# ════════════════════════════════════════════════

def _detecter_nb_colonnes(blocs: list, largeur: float) -> int:
    """
    Détecte automatiquement 1, 2 ou 3 colonnes
    via KMeans sur les centres X des blocs.
    Universel — fonctionne pour tout format de CV.
    """
    if len(blocs) < 4:
        return 1

    centres_x = np.array([
        ((b[0] + b[2]) / 2) / largeur
        for b in blocs
    ]).reshape(-1, 1)

    meilleur_k     = 1
    meilleur_score = -1

    for k in range(2, 4):
        if len(centres_x) <= k * 2:
            break
        try:
            km     = KMeans(n_clusters=k, random_state=42, n_init=10)
            labels = km.fit_predict(centres_x)
            score  = silhouette_score(centres_x, labels)
            if score > 0.42 and score > meilleur_score:
                meilleur_score = score
                meilleur_k     = k
        except Exception:
            pass

    return meilleur_k


def _trier_blocs_par_colonnes(blocs: list, nb_colonnes: int,
                               largeur: float) -> list:
    """
    Trie les blocs colonne par colonne puis par Y dans chaque colonne.
    Garantit lecture naturelle gauche→droite, haut→bas par colonne.
    """
    if nb_colonnes == 1:
        return sorted(blocs, key=lambda b: b[1])

    largeur_col = largeur / nb_colonnes

    def cle_tri(bloc):
        centre_x = (bloc[0] + bloc[2]) / 2
        col      = min(int(centre_x // largeur_col), nb_colonnes - 1)
        return (col, bloc[1])

    return sorted(blocs, key=cle_tri)


# ════════════════════════════════════════════════
# 3. PDF NATIF — pymupdf + positions
# ════════════════════════════════════════════════

def _extraire_pdf_natif(path: str):
    """
    Extrait texte + positions X/Y depuis PDF natif.
    Retourne : (texte_brut: str, blocs_pos: list)
    """
    doc        = fitz.open(path)
    text_final = []
    blocs_pos  = []

    for page in doc:
        largeur     = page.rect.width
        blocs_bruts = page.get_text("blocks")

        # Garder uniquement blocs texte non vides
        blocs_texte = [
            b for b in blocs_bruts
            if b[6] == 0 and b[4].strip()
        ]

        if not blocs_texte:
            continue

        # Détecter colonnes + trier
        nb_colonnes = _detecter_nb_colonnes(blocs_texte, largeur)
        blocs_tries = _trier_blocs_par_colonnes(blocs_texte, nb_colonnes, largeur)

        for bloc in blocs_tries:
            # pymupdf regroupe les phrases multi-lignes → replace \n par espace
            texte = bloc[4].strip().replace("\n", " ")
            texte = re.sub(r' {2,}', ' ', texte)
            if not texte:
                continue

            centre_x = (bloc[0] + bloc[2]) / 2
            col      = min(
                int(centre_x / (largeur / max(nb_colonnes, 1))),
                nb_colonnes - 1
            )

            text_final.append(texte)
            blocs_pos.append({
                "texte":       texte,
                "col":         col,
                "y":           bloc[1],
                "nb_colonnes": nb_colonnes
            })

    doc.close()
    return "\n".join(text_final), blocs_pos


# ════════════════════════════════════════════════
# 4. PDF SCANNÉ — tesseract OCR
# ════════════════════════════════════════════════

def _detecter_nb_colonnes_ocr(centres_x: np.ndarray, largeur: float) -> int:
    if len(centres_x) < 4:
        return 1

    x_norm     = (centres_x / largeur).reshape(-1, 1)
    meilleur_k = 1
    meilleur_s = -1

    for k in range(2, 4):
        if len(x_norm) <= k * 2:
            break
        try:
            km     = KMeans(n_clusters=k, random_state=42, n_init=10)
            labels = km.fit_predict(x_norm)
            score  = silhouette_score(x_norm, labels)
            if score > 0.42 and score > meilleur_s:
                meilleur_s = score
                meilleur_k = k
        except Exception:
            pass

    return meilleur_k


def _extraire_pdf_scanne(path: str):
    """
    Extrait texte depuis PDF scanné via tesseract OCR.
    Retourne : (texte_brut: str, blocs_pos: list)
    """
    pages      = convert_from_path(path, dpi=300)
    text_final = []
    blocs_pos  = []

    for num_page, image in enumerate(pages):
        largeur_image = image.width

        data = pytesseract.image_to_data(
            image,
            output_type=pytesseract.Output.DICT,
            lang="fra+eng",
            config="--psm 11"
        )

        # Construire blocs filtrés
        mots = []
        for i in range(len(data["text"])):
            mot = data["text"][i].strip()
            try:
                conf = int(data["conf"][i])
            except Exception:
                conf = 0
            if conf >= 50 and mot:
                mots.append({
                    "text": mot,
                    "x":    data["left"][i],
                    "y":    data["top"][i],
                    "w":    data["width"][i],
                    "h":    data["height"][i],
                })

        if not mots:
            continue

        # Détecter colonnes
        centres_x   = np.array([m["x"] + m["w"] / 2 for m in mots])
        nb_colonnes = _detecter_nb_colonnes_ocr(centres_x, largeur_image)
        largeur_col = largeur_image / nb_colonnes

        # Trier par colonne puis Y
        mots_tries = sorted(mots, key=lambda m: (
            min(int((m["x"] + m["w"] / 2) // largeur_col), nb_colonnes - 1),
            m["y"]
        ))

        # Reconstruire lignes cohérentes
        lignes   = []
        courante = []
        y_prev   = None
        SEUIL_Y  = 10

        for m in mots_tries:
            if y_prev is None or abs(m["y"] - y_prev) <= SEUIL_Y:
                courante.append(m["text"])
            else:
                if courante:
                    ligne = " ".join(courante).strip()
                    if ligne:
                        lignes.append(ligne)
                        col = min(
                            int((mots_tries[0]["x"] + mots_tries[0]["w"] / 2) // largeur_col),
                            nb_colonnes - 1
                        )
                        blocs_pos.append({
                            "texte":       ligne,
                            "col":         col,
                            "y":           y_prev,
                            "nb_colonnes": nb_colonnes
                        })
                courante = [m["text"]]
            y_prev = m["y"]

        if courante:
            ligne = " ".join(courante).strip()
            if ligne:
                lignes.append(ligne)
                blocs_pos.append({
                    "texte":       ligne,
                    "col":         0,
                    "y":           y_prev or 0,
                    "nb_colonnes": nb_colonnes
                })

        text_final.extend(lignes)

    return "\n".join(text_final), blocs_pos


# ════════════════════════════════════════════════
# 5. ROUTER PDF
# ════════════════════════════════════════════════

def _extraire_pdf(path: str):
    """
    Détecte automatiquement PDF natif ou scanné.
    Natif  → pymupdf (précis, rapide)
    Scanné → tesseract OCR (fallback)
    """
    doc        = fitz.open(path)
    texte_test = ""

    for page in doc:
        texte_test += page.get_text("text")
        if len(texte_test.strip()) > 150:
            break

    doc.close()

    if len(texte_test.strip()) > 150:
        return _extraire_pdf_natif(path)
    else:
        return _extraire_pdf_scanne(path)


# ════════════════════════════════════════════════
# 6. NETTOYAGE
# ════════════════════════════════════════════════

def _nettoyer(text: str) -> str:
    if not text:
        return ""

    # Corriger erreurs OCR lettre+chiffre collés
    text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)
    text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', text)

    # Corriger email cassé par OCR
    text = re.sub(r'(\S+@\S+)\s*\.\s*(com|fr|ma|org)', r'\1.\2', text)

    # Nettoyer espaces multiples
    text = re.sub(r' {2,}', ' ', text)

    # Nettoyer lignes vides multiples
    lignes = [l.strip() for l in text.splitlines()]
    propre           = []
    vide_precedente  = False

    for ligne in lignes:
        if not ligne:
            if not vide_precedente:
                propre.append("")
            vide_precedente = True
        else:
            if len(ligne) >= 2:
                propre.append(ligne)
            vide_precedente = False

    return "\n".join(propre).strip()


# ════════════════════════════════════════════════
# 7. NETTOYAGE OFFRE LINKEDIN
# ════════════════════════════════════════════════

def _nettoyer_offre(text: str) -> str:
    """
    Nettoyage spécifique pour offres LinkedIn copiées-collées.
    Gère : emojis, bullets, lignes vides multiples, caractères parasites.
    """
    if not text:
        return ""

    # Supprimer emojis et caractères spéciaux LinkedIn
    text = re.sub(r'[^\w\s\-\.:/@,;\'\"()+°#&éèêëàâùûôîïç\n]', ' ', text)

    # Normaliser bullets LinkedIn (•, -, *, →)
    text = re.sub(r'^\s*[•\-\*→✓✔►▪]\s*', '', text, flags=re.MULTILINE)

    # Nettoyer espaces multiples
    text = re.sub(r' {2,}', ' ', text)

    # Nettoyer lignes vides multiples
    lignes = [l.strip() for l in text.splitlines()]
    propre          = []
    vide_precedente = False

    for ligne in lignes:
        if not ligne:
            if not vide_precedente:
                propre.append("")
            vide_precedente = True
        else:
            if len(ligne) >= 2:
                propre.append(ligne)
            vide_precedente = False

    return "\n".join(propre).strip()


# ════════════════════════════════════════════════
# 8. API PUBLIQUE
# ════════════════════════════════════════════════

def extraire_cv(path: str) -> dict:
    """
    Extrait le texte d'un CV (PDF ou DOCX).

    Retourne :
    {
        "type"     : "cv",
        "content"  : "texte brut nettoyé",
        "blocs_pos": [{"texte", "col", "y", "nb_colonnes"}, ...]
    }
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        text, blocs_pos = _extraire_pdf(path)
    elif ext == ".docx":
        text      = _extraire_docx(path)
        blocs_pos = []
    else:
        raise ValueError(f"Format non supporté : {ext} — utilisez PDF ou DOCX")

    return {
        "type":      "cv",
        "content":   _nettoyer(text),
        "blocs_pos": blocs_pos
    }


def extraire_offre(json_offre: str) -> dict:
    """
    Parse et nettoie une offre LinkedIn reçue en JSON.

    Entrée  : '{"offre": "texte brut LinkedIn..."}'
    Retourne:
    {
        "type"    : "offre",
        "content" : "texte nettoyé"
    }
    """
    try:
        data = json.loads(json_offre)
    except json.JSONDecodeError:
        raise ValueError("JSON offre invalide — format attendu : {\"offre\": \"...\"}")

    text = data.get("offre", "").strip()

    if not text:
        raise ValueError("Clé 'offre' vide ou manquante dans le JSON")

    return {
        "type":    "offre",
        "content": _nettoyer_offre(text)
    }